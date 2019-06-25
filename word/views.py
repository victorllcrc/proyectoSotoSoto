from django.shortcuts import render, HttpResponse
from django.core.files.storage import FileSystemStorage
import pandas as pd
from docx import Document
from renderword.settings import BASE_DIR
#import docx
import json
# Create your views here.



def home(request):
    return render(request, 'word/base.html')


def cargaexcel(request):
    if request.method == 'POST':
        uploaded_file = request.FILES['document']
        fs = FileSystemStorage()
        name = fs.save(uploaded_file.name, uploaded_file)

        xls = pd.ExcelFile(BASE_DIR + fs.url(name))
        hojas = xls.sheet_names
        print(len(hojas))
        dc1 = Document()
        def mat(abcd):
            pri = 0
            val = 0
            aux = []

            for i in range(len(abcd)):
                aux.append([])
                for j in range(len(abcd[i])):
                    aux[i].append(None)

            for i in range(len(abcd)):
                sec = 0
                val = 0
                for j in range(len(abcd[i])):

                    if type(abcd[i][j]) is str:
                        val += 1
                        aux[pri][sec] = abcd[i][j]
                        sec += 1
                    else:
                        if type(abcd[i][j]) is int:
                            abcd[i][j] = str(abcd[i][j])
                            val += 1
                            aux[pri][sec] = abcd[i][j]
                            sec += 1

                if val > 0:
                    pri += 1

            return aux

        for hoja in hojas:

            df = xls.parse(hoja)

            fich = df.__array__()

            real = mat(fich);
            fich = real
            print(fich)



            tit = fich[0][0]
            anio = str(fich[4][1])
            nota = str(fich[10][1])

            dc1.add_heading(tit, 0)
            dc1.add_heading(fich[1][0], 2)
            p = dc1.add_paragraph()
            p.add_run('\n')
            p.add_run(fich[2][0] + ' \t').bold = True
            p.add_run(fich[2][1])

            p.add_run('\n\n')
            p.add_run(fich[3][0] + ' \t').bold = True
            p.add_run(fich[3][1])

            p.add_run('\n\n')
            p.add_run(fich[4][0] + ' \t').bold = True
            p.add_run(anio)

            dc1.add_heading(fich[5][0], 2)
            p1 = dc1.add_paragraph()
            p1.add_run('\n')
            p1.add_run(fich[6][0] + ' \t').bold = True
            p1.add_run(fich[6][1])

            p1.add_run('\n\n')
            p1.add_run(fich[7][0] + ' \t').bold = True
            p1.add_run(fich[7][1])

            dc1.add_heading(fich[8][0], 2)
            p2 = dc1.add_paragraph()
            p2.add_run('\n')
            p2.add_run(fich[9][0] + ' \t').bold = True
            p2.add_run(fich[9][1])

            p2.add_run('\n\n')
            p2.add_run(fich[10][0] + ' \t').bold = True
            p2.add_run(nota)

            dc1.add_heading(fich[11][0], 2)
            p3 = dc1.add_paragraph()
            p3.add_run('\n')
            p3.add_run(fich[12][0] + ' \t').bold = True
            p3.add_run(fich[12][1])

            dc1.add_page_break()


        nombre_archivo = "Proyecto.docx"
        response = HttpResponse(content_type="application/msword")
        contenido = "attachment; filename= {0}".format(nombre_archivo)
        response["Content-Disposition"] = contenido
        dc1.save(response)

        return response
        #return render(request, 'word/index.html', {'url':fs.url(name)})

    return render(request, 'word/index.html')

